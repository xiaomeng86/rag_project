from __future__ import annotations

from io import BytesIO
import json
from pathlib import Path

from docx import Document
import pdfplumber
import redis

from config import (
    REDIS_DB,
    REDIS_HOST,
    REDIS_PORT,
    SUPPORTED_TEMPORARY_EXTENSIONS,
    TEMPORARY_DOCUMENT_TTL_SECONDS,
)


def _redis_client() -> redis.Redis:
    return redis.Redis(
        host=REDIS_HOST,
        port=REDIS_PORT,
        db=REDIS_DB,
        decode_responses=True,
    )


def _key(session_id: str) -> str:
    return f"session:{session_id}:temporary_document"


def _parse_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                parts.append(text)
    return "\n".join(parts)


def _parse_pdf(content: bytes) -> str:
    with pdfplumber.open(BytesIO(content)) as pdf:
        return "\n".join((page.extract_text() or "").strip() for page in pdf.pages).strip()


def _parse_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("TXT 编码无法识别")


def parse_temporary_document(file_name: str, content: bytes) -> tuple[str, str]:
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_TEMPORARY_EXTENSIONS:
        raise ValueError("临时文档仅支持 PDF、DOCX、TXT")
    if suffix == ".pdf":
        text = _parse_pdf(content)
    elif suffix == ".docx":
        text = _parse_docx(content)
    else:
        text = _parse_txt(content)
    if not text:
        raise ValueError("文档中没有可解析文本")
    return suffix.removeprefix("."), text


def store_temporary_document(session_id: str, file_name: str, content: bytes) -> dict:
    document_type, text = parse_temporary_document(file_name, content)
    record = {
        "document_id": f"temporary-{session_id}",
        "document_name": Path(file_name).name,
        "document_type": document_type,
        "file_size": len(content),
        "content": text,
    }
    _redis_client().setex(
        _key(session_id),
        TEMPORARY_DOCUMENT_TTL_SECONDS,
        json.dumps(record, ensure_ascii=False),
    )
    return {**record, "expires_in_seconds": TEMPORARY_DOCUMENT_TTL_SECONDS}


def get_temporary_document(session_id: str) -> dict | None:
    client = _redis_client()
    value = client.get(_key(session_id))
    if not value:
        return None
    record = json.loads(value)
    record["expires_in_seconds"] = max(client.ttl(_key(session_id)), 0)
    return record


def redis_health() -> bool:
    return bool(_redis_client().ping())

