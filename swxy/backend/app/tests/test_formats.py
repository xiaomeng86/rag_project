from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from config import SUPPORTED_KNOWLEDGE_EXTENSIONS, SUPPORTED_TEMPORARY_EXTENSIONS
from service.core import file_parse
from service.core.rag.app.naive import chunk
from service.temporary_document_service import parse_temporary_document


def _callback(*_args, **_kwargs):
    return None


def _pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()
    document = canvas.Canvas(buffer)
    document.drawString(72, 720, text)
    document.save()
    return buffer.getvalue()


def test_supported_format_contract_is_exact():
    assert SUPPORTED_KNOWLEDGE_EXTENSIONS == {
        ".pdf", ".docx", ".xlsx", ".txt", ".md", ".markdown", ".html", ".htm"
    }
    assert SUPPORTED_TEMPORARY_EXTENSIONS == {".pdf", ".docx", ".txt"}


def test_six_permanent_formats_use_runtime_generated_documents(tmp_path: Path):
    text = "GSK synthetic document explains enterprise retrieval evidence."
    paths: list[tuple[Path, dict]] = []

    txt_path = tmp_path / "sample.txt"
    txt_path.write_text(text, encoding="utf-8")
    paths.append((txt_path, {}))

    markdown_path = tmp_path / "sample.md"
    markdown_path.write_text(f"# Evidence\n\n{text}", encoding="utf-8")
    paths.append((markdown_path, {}))

    html_path = tmp_path / "sample.html"
    html_path.write_text(f"<html><body><p>{text}</p></body></html>", encoding="utf-8")
    paths.append((html_path, {}))

    docx_path = tmp_path / "sample.docx"
    docx = Document()
    docx.add_paragraph(text)
    docx.save(docx_path)
    paths.append((docx_path, {}))

    xlsx_path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    workbook.active.append(["topic", "content"])
    workbook.active.append(["RAG", text])
    workbook.save(xlsx_path)
    paths.append((xlsx_path, {}))

    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(_pdf_bytes(text))
    paths.append((pdf_path, {"parser_config": {"chunk_token_num": 128, "layout_recognize": "Plain Text"}}))

    for path, kwargs in paths:
        parsed = chunk(str(path), callback=_callback, **kwargs)
        assert parsed, path.suffix
        assert any(item.get("content_with_weight") for item in parsed)


def test_default_pdf_route_selects_deepdoc(monkeypatch, tmp_path: Path):
    from service.core.rag.app import naive

    called = {"deepdoc": False}

    class FakePdf:
        def __call__(self, *_args, **_kwargs):
            called["deepdoc"] = True
            return [("DeepDoc route", "")], []

        def crop(self, *_args, **_kwargs):
            raise NotImplementedError

    monkeypatch.setattr(naive, "Pdf", FakePdf)
    path = tmp_path / "route.pdf"
    path.write_bytes(_pdf_bytes("route"))
    assert chunk(str(path), callback=_callback)
    assert called["deepdoc"] is True


def test_ingestion_auto_routes_pdf_by_text_layer(monkeypatch, tmp_path: Path):
    selected_modes: list[str] = []

    def fake_chunk(_path, **kwargs):
        selected_modes.append(kwargs["parser_config"]["layout_recognize"])
        return [{"content_with_weight": "evidence"}]

    monkeypatch.setattr(file_parse, "chunk", fake_chunk)

    searchable = tmp_path / "searchable.pdf"
    searchable.write_bytes(_pdf_bytes("Searchable PDF evidence contains enough text."))
    scanned = tmp_path / "scanned.pdf"
    scanned.write_bytes(_pdf_bytes(""))

    assert file_parse.parse(searchable)
    assert file_parse.parse(scanned)
    assert selected_modes == ["Plain Text", "DeepDOC"]


def test_128_tokens_is_a_chunking_target(tmp_path: Path):
    path = tmp_path / "long.txt"
    path.write_text("\n".join([f"Evidence sentence {index}." for index in range(300)]), encoding="utf-8")
    parsed = chunk(str(path), callback=_callback)
    assert len(parsed) > 1


def test_three_temporary_formats_are_parsed_from_memory():
    document = Document()
    document.add_paragraph("temporary DOCX evidence")
    docx_buffer = BytesIO()
    document.save(docx_buffer)

    cases = [
        ("temporary.txt", "temporary TXT evidence".encode(), "txt"),
        ("temporary.docx", docx_buffer.getvalue(), "docx"),
        ("temporary.pdf", _pdf_bytes("temporary PDF evidence"), "pdf"),
    ]
    for file_name, content, expected_type in cases:
        document_type, text = parse_temporary_document(file_name, content)
        assert document_type == expected_type
        assert "temporary" in text.lower()
