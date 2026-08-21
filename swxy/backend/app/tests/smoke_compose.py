"""End-to-end acceptance test for a running Compose stack using synthetic files."""

from __future__ import annotations

import argparse
from io import BytesIO
import json
import time

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
import requests


REQUEST_TIMEOUT = (10, 600)
EVIDENCE = "GSK POC synthetic evidence confirms the closed loop works in 2026."


def _docx_bytes() -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_heading("Synthetic DOCX", level=1)
    document.add_paragraph(EVIDENCE)
    document.save(stream)
    return stream.getvalue()


def _xlsx_bytes() -> bytes:
    stream = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Evidence"
    sheet.append(["source", "content"])
    sheet.append(["synthetic", EVIDENCE])
    workbook.save(stream)
    return stream.getvalue()


def _pdf_bytes() -> bytes:
    stream = BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.setTitle("Synthetic PDF")
    pdf.drawString(72, 760, "GSK POC synthetic PDF evidence")
    pdf.drawString(72, 740, "The closed loop works in 2026.")
    pdf.showPage()
    pdf.save()
    return stream.getvalue()


def _expect(response: requests.Response, status: int) -> dict | list:
    if response.status_code != status:
        raise AssertionError(
            f"{response.request.method} {response.url}: "
            f"expected {status}, got {response.status_code}: {response.text[:1000]}"
        )
    return response.json()


def _upload(base_url: str, headers: dict[str, str], name: str, content: bytes) -> dict:
    response = requests.post(
        f"{base_url}/api/v1/knowledge/documents",
        headers=headers,
        files={"file": (name, content, "application/octet-stream")},
        timeout=REQUEST_TIMEOUT,
    )
    value = _expect(response, 201)
    assert isinstance(value, dict)
    assert value["status"] == "ready"
    assert value["chunk_count"] > 0
    return value


def _put_temporary(
    base_url: str,
    headers: dict[str, str],
    session_id: str,
    name: str,
    content: bytes,
) -> dict:
    response = requests.put(
        f"{base_url}/api/v1/sessions/{session_id}/temporary-document",
        headers=headers,
        files={"file": (name, content, "application/octet-stream")},
        timeout=REQUEST_TIMEOUT,
    )
    value = _expect(response, 200)
    assert isinstance(value, dict)
    assert value["expires_in_seconds"] == 7200
    current = _expect(
        requests.get(
            f"{base_url}/api/v1/sessions/{session_id}/documents",
            headers=headers,
            timeout=REQUEST_TIMEOUT,
        ),
        200,
    )
    assert isinstance(current, dict)
    assert current["has_document"] is True
    assert current["document"]["document_name"] == name
    return value


def _chat(
    base_url: str,
    headers: dict[str, str],
    session_id: str,
) -> tuple[list[str], dict[str, dict]]:
    response = requests.post(
        f"{base_url}/api/v1/sessions/{session_id}/chat",
        headers=headers,
        json={"message": "What synthetic evidence confirms the closed loop?"},
        stream=True,
        timeout=REQUEST_TIMEOUT,
    )
    if response.status_code != 200:
        raise AssertionError(f"chat returned {response.status_code}: {response.text[:1000]}")

    event_name = ""
    event_order: list[str] = []
    last_payload: dict[str, dict] = {}
    for line in response.iter_lines(decode_unicode=True):
        if not line:
            continue
        if line.startswith("event: "):
            event_name = line.removeprefix("event: ")
        elif line.startswith("data: "):
            payload = json.loads(line.removeprefix("data: "))
            assert payload["type"] == event_name
            event_order.append(event_name)
            last_payload[event_name] = payload

    required = {"citations", "thinking", "answer", "recommendations", "done"}
    assert required.issubset(event_order)
    assert "error" not in event_order
    assert event_order[0] == "citations"
    citations = last_payload["citations"]["citations"]
    assert citations and citations[0]["citation_id"] == 1
    assert citations[0]["chunk_id"]
    assert citations[0]["document_id"]
    assert citations[0]["document_name"]
    assert citations[0]["content"]
    return event_order, last_payload


def run(base_url: str) -> dict:
    username = f"compose-smoke-{int(time.time())}"
    credentials = {"username": username, "password": "synthetic-pass-123"}
    _expect(
        requests.post(
            f"{base_url}/api/v1/auth/register",
            json=credentials,
            timeout=REQUEST_TIMEOUT,
        ),
        201,
    )
    login = _expect(
        requests.post(
            f"{base_url}/api/v1/auth/login",
            json=credentials,
            timeout=REQUEST_TIMEOUT,
        ),
        200,
    )
    assert isinstance(login, dict)
    headers = {"Authorization": f"Bearer {login['access_token']}"}

    created_session = _expect(
        requests.post(
            f"{base_url}/api/v1/sessions",
            headers=headers,
            timeout=REQUEST_TIMEOUT,
        ),
        201,
    )
    assert isinstance(created_session, dict)
    session_id = created_session["session_id"]
    sessions = _expect(
        requests.get(
            f"{base_url}/api/v1/sessions",
            headers=headers,
            timeout=REQUEST_TIMEOUT,
        ),
        200,
    )
    assert isinstance(sessions, dict)
    assert any(item["session_id"] == session_id for item in sessions["sessions"])

    permanent_files = {
        "synthetic.txt": EVIDENCE.encode(),
        "synthetic.md": f"# Synthetic Markdown\n\n{EVIDENCE}".encode(),
        "synthetic.html": f"<h1>Synthetic HTML</h1><p>{EVIDENCE}</p>".encode(),
        "synthetic.docx": _docx_bytes(),
        "synthetic.xlsx": _xlsx_bytes(),
        "synthetic.pdf": _pdf_bytes(),
    }
    documents = [
        _upload(base_url, headers, name, content)
        for name, content in permanent_files.items()
    ]
    listed = _expect(
        requests.get(
            f"{base_url}/api/v1/knowledge/documents",
            headers=headers,
            timeout=REQUEST_TIMEOUT,
        ),
        200,
    )
    assert isinstance(listed, list) and len(listed) == len(permanent_files)

    _put_temporary(
        base_url,
        headers,
        session_id,
        "temporary.txt",
        EVIDENCE.encode(),
    )
    _put_temporary(
        base_url,
        headers,
        session_id,
        "temporary.docx",
        _docx_bytes(),
    )
    _put_temporary(
        base_url,
        headers,
        session_id,
        "temporary.pdf",
        _pdf_bytes(),
    )

    event_order, _ = _chat(base_url, headers, session_id)
    history = _expect(
        requests.get(
            f"{base_url}/api/v1/sessions/{session_id}/messages",
            headers=headers,
            timeout=REQUEST_TIMEOUT,
        ),
        200,
    )
    assert isinstance(history, list) and len(history) == 1
    assert history[0]["citations"][0]["citation_id"] == 1
    assert history[0]["recommendations"]

    for document in documents:
        deleted = _expect(
            requests.delete(
                f"{base_url}/api/v1/knowledge/documents/{document['document_id']}",
                headers=headers,
                timeout=REQUEST_TIMEOUT,
            ),
            200,
        )
        assert isinstance(deleted, dict)
        assert deleted["deleted_chunks"] == document["chunk_count"]
    assert _expect(
        requests.get(
            f"{base_url}/api/v1/knowledge/documents",
            headers=headers,
            timeout=REQUEST_TIMEOUT,
        ),
        200,
    ) == []

    return {
        "permanent_formats": len(permanent_files),
        "temporary_formats": 3,
        "sse_events": event_order,
        "history_messages": len(history),
        "deleted_documents": len(documents),
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--base-url", default="http://host.docker.internal:8000")
    arguments = parser.parse_args()
    print(json.dumps(run(arguments.base_url.rstrip("/")), ensure_ascii=False))
